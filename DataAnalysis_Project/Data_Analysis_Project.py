#Divine Lotanna Mbamara

"""
DESCRIPTION:
This program reads an Excel spreadsheet containing student performance data, cleans and organizes the data, conducts analysis, and generates
visualizations to highlight key patterns and insights.

*** Data Cleaning
 1. Student Column: Renamed 'student' column to 'Student' & Standardized column data to ensure consistent capitalization (Title case)
 2. Study Hours Column: Renamed 'studyhours' column to 'Study Hours' & Deleted rows containing outliers (negative values) in the Study Hours column
 3. Attendance Rate Column: Renamed 'attendance_rate' column to 'Attendance Rate in %'
 4. Homework Column: Renamed 'homework' column to 'Homework (20 marks)'
 5. Participation Score Column: Renamed 'participationscore' column to 'Participation Score (15 marks)'
 6. Midterm Score Column: Renamed 'previousscore' column to 'Midterm Score (30 marks)'
 7. Final Exam Score Column: Renamed 'final_exam_score' column to 'Final Exam Score (35 marks)'
 8. Status Column: Renamed 'pass_fail' column to 'Status', Modified 0s to Fail & 1s to Pass
 9. Duplicates: Removed duplicates
 10. New Total Score Column: Created a 'Total Score' column to hold the sum of Homework, Participation, Midterm, and Final Exam
 11. Save: Saved the cleaned student performance excel spreadsheet to the same folder as the original spreadsheet 

*** Data Analysis & Data Visualization
 1. Set Up Report: Opened a Word document and added a report title.
 2. Loaded Data: Read the cleaned Excel file with student performance data.
 3. Overall Statistics: Calculated summary statistics (using describe()) for the Total Score and added them to the word document.
 4. Histogram: Created and saved a histogram (with 20 bins and a density curve) to show the distribution of Total Scores, then inserted it
    into the word document.
 5. Boxplot (Total Score): Generated a boxplot comparing Total Scores between Pass and Fail groups, saved, and added it to the report.
 6. Countplot: Made a countplot (using pastel colors) to show the distribution of Pass/Fail statuses, saved and added it to the report.
 7. Study Hours Correlation: Computed the correlation between Study Hours and Total Score and added that number to the word document.
 8. Scatter Plot: Created a scatterplot showing the relationship between Study Hours and Total Score, then saved and inserted it into the word document.
 9. Boxplot (Study Hours): Generated a boxplot to compare Study Hours by Pass/Fail status, saved it, and inserted it.
 10. Participation Score Correlation: Calculated the correlation between Participation Score and Total Score and added it to the report.
 11. Boxplot (Participation Score): Created a boxplot comparing Participation Scores by Pass/Fail status, saved and inserted it into the word document.
 12. Pivot Table: Built a pivot table showing the average Total Score for each Study Hour value grouped by Status, and added it to the document.
 13. Final Save: Saved the Word document containing all the analysis and visualizations.

"""


#Import modules
import pandas as pd
import tkinter as tk
from tkinter import filedialog
import os
import matplotlib.pyplot as plt
import seaborn as sns
import docx


# Initialize Tkinter
root = tk.Tk()
root.withdraw()  # Hide the root window

original_file_path = None

# Use dialog box to select file. Keep asking until a valid Excel file is selected
while not original_file_path:
    original_file_path = filedialog.askopenfilename(title="Select the Student Performance File",
                                           filetypes=[("Excel Files", "*.xlsx *.xls")])
    if not original_file_path:
        print("No file selected. Please select an Excel file:")


# Read the selected file into pandas and load the spreadsheet
studentData = pd.read_excel(original_file_path)


#DATA CLEANING
# Create a backup (copy of the raw file)
studentData.to_excel('backup_raw_student_performance.xlsx', index=False)

# Correction 1: Rename 'student' column & Clean the data in the column for consistent capitalization (Title case)
studentData.rename(columns={'student': 'Student'}, inplace=True)
studentData['Student'] = studentData['Student'].str.title()

# Correction 2: Rename 'studyhours' column & Delete outliers (negative values) in the column
studentData.rename(columns={'studyhours': 'Study Hours'}, inplace=True)
studentData = studentData[studentData['Study Hours'] >= 0]

# Correction 3: Rename 'attendance_rate' column (add % sign)
studentData.rename(columns={'attendance_rate': 'Attendance Rate in %'}, inplace=True)

# Correction 4: Rename 'homework' column
studentData.rename(columns={'homework': 'Homework (20 marks)'}, inplace=True)

# Correction 5: Rename 'participationscore' column
studentData.rename(columns={'participationscore': 'Participation Score (15 marks)'}, inplace=True)

# Correction 6: Rename 'previousscore' column
studentData.rename(columns={'previous_score': 'Midterm Score (30 marks)'}, inplace=True)

# Correction 7: Rename 'final_exam_score' column
studentData.rename(columns={'final_exam_score': 'Final Exam Score (35 marks)'}, inplace=True)

# Correction 8: Fix the pass_fail column - Rename column name to 'Status', Change 0s to Fail & 1s to Pass
studentData.rename(columns={'pass_fail': 'Status'}, inplace=True)
studentData['Status'] = studentData['Status'].map({0: 'Fail', 1: 'Pass'})

# Correction 9: Remove duplicates
studentData = studentData.drop_duplicates()

# Correction 10: Create a 'Total Score' column as the sum of Homework, Participation, Midterm, and Final Exam
studentData['Total Score'] = studentData[['Homework (20 marks)', 'Participation Score (15 marks)', 'Midterm Score (30 marks)', 'Final Exam Score (35 marks)']].sum(axis=1)

#Save
#Create a name for the cleaned file
cleaned_file_name = "cleaned_student_performance.xlsx"

# Get the folder path where the original file is located
folder_path = os.path.dirname(original_file_path)

# Create the full path for the cleaned file
cleaned_file_path = os.path.join(folder_path, cleaned_file_name)

# Save the cleaned data
studentData.to_excel(cleaned_file_path, index=False)

# Print the save location
print(f"Cleaned data saved in: {cleaned_file_path}")



# DATA ANALYSIS & VISUALIZATION
#store data analysis and visualization in a word doc
doc = docx.Document("student_performance_analysis_visualization.docx")
doc.add_heading("Student Performance Analysis and Visualization Report", 1)

# Load cleaned student data
studentData = pd.read_excel(cleaned_file_path)

# Calculate overall performance statistics
overall_performance = studentData['Total Score'].describe()
doc.add_heading("Overall Performance Statistics", 3)
doc.add_paragraph(str(overall_performance.round(2)))
doc.add_paragraph("")

# Histogram for overall performance (Total Score)
plt.figure(figsize=(8,6))
sns.histplot(studentData['Total Score'], bins=20, kde=True, color='blue')
plt.title('Distribution of Total Scores')
plt.xlabel('Total Score')
plt.ylabel('Frequency')
plt.savefig("overall_performance_histogram.png")
doc.add_heading("Histogram Displaying Distribution of Total Scores", 3)
doc.add_picture("overall_performance_histogram.png", width=docx.shared.Inches(6), height=docx.shared.Cm(8))
doc.add_page_break()

# Boxplot to compare Total Scores for Pass/Fail
plt.figure(figsize=(8,6))
sns.boxplot(x='Status', y='Total Score', data=studentData)
plt.title('Total Scores by Pass/Fail Status')
plt.xlabel('Status')
plt.ylabel('Total Score')
plt.savefig("total_scores_pass-fail_boxplot.png")
doc.add_heading("Boxplot Comparing Total Scores by Pass/Fail Status", 3)
doc.add_picture("total_scores_pass-fail_boxplot.png", width=docx.shared.Inches(6), height=docx.shared.Cm(8))
doc.add_page_break()

# Plot Pass/Fail Distribution
plt.figure(figsize=(6, 4))
sns.countplot(x="Status", data=studentData, hue="Status", palette="pastel", legend=False)
plt.title("Pass/Fail Distribution")
plt.xlabel("Status (Pass/Fail)")
plt.ylabel("Count")
plt.savefig("pass-fail_distribution_countplot.png")
doc.add_heading("Countplot Displaying Pass/Fail Distribution", 3)
doc.add_picture("pass-fail_distribution_countplot.png", width=docx.shared.Inches(6), height=docx.shared.Cm(8))
doc.add_paragraph("")

# Correlation between Study Hours and Total Score
study_hours_corr = studentData['Study Hours'].corr(studentData['Total Score'])
doc.add_heading("Correlation between Study Hours and Total Score", 3)
doc.add_paragraph("Correlation between Study Hours and Total Score: " + str(study_hours_corr.round(2)))
doc.add_page_break()

# Scatter Plot: Study Hours vs Total Score
plt.figure(figsize=(8, 5))
sns.scatterplot(x=studentData["Study Hours"], y=studentData["Total Score"])
plt.title("Study Hours vs. Total Score")
plt.xlabel("Study Hours")
plt.ylabel("Total Score (%)")
plt.savefig("study-hours_vs_total-score_scatterplot.png")
doc.add_heading("Scatterplot Displaying the Relationship between Study Hours and Total Score", 3)
doc.add_picture("study-hours_vs_total-score_scatterplot.png", width=docx.shared.Inches(6), height=docx.shared.Cm(8))
doc.add_page_break()

# Compare Study Hours for Pass/Fail students
plt.figure(figsize=(8,6))
sns.boxplot(x='Status', y='Study Hours', data=studentData)
plt.title('Study Hours by Pass/Fail Status')
plt.xlabel('Status')
plt.ylabel('Study Hours')
plt.savefig("study-hours_vs_pass-fail-status_boxplot.png")
doc.add_heading("Boxplot Displaying the Relationship between Study Hours and Pass/Fail Status", 3)
doc.add_picture("study-hours_vs_pass-fail-status_boxplot.png", width=docx.shared.Inches(6), height=docx.shared.Cm(8))
doc.add_paragraph("")

# Correlation between Participation Score and Total Score
participation_score_corr = studentData['Participation Score (15 marks)'].corr(studentData['Total Score'])
doc.add_heading("Correlation between Participation Score and Total Score", 3)
doc.add_paragraph("The correlation between participation score and total Score is " + str(participation_score_corr.round(2)))
doc.add_page_break()

# Compare Participation Scores for Pass/Fail students
plt.figure(figsize=(8,6))
sns.boxplot(x='Status', y='Participation Score (15 marks)', data=studentData)
plt.title('Participation Score by Pass/Fail Status')
plt.xlabel('Status')
plt.ylabel('Participation Score')
plt.savefig("participation-scores_vs_pass-fail-status_boxplot.png")
doc.add_heading("Boxplot Displaying the Relationship between Participation Scores and Pass/Fail Status", 3)
doc.add_picture("participation-scores_vs_pass-fail-status_boxplot.png", width=docx.shared.Inches(6), height=docx.shared.Cm(8))
doc.add_paragraph("")

# Create a pivot table for Study Hours vs Total Score by Pass/Fail
pivot_study_score = studentData.pivot_table(values='Total Score', index='Status', columns='Study Hours', aggfunc='mean')
doc.add_heading("Pivot Table Showing Study Hours vs Total Score by Pass/Fail", 3)
doc.add_paragraph(str(pivot_study_score.round(2)))

#save the word doc
doc.save("student_performance_analysis_visualization.docx")




